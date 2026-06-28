"""
AGL - Contract Writer
Genere le DOCX redline final a partir du resultat de negociation.
Rouge = position Buyer | Vert = position Seller | Bleu = accord
"""

import logging
import os
from datetime import datetime
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.shared import Cm, Inches

from src.contracts.nda_template import (
    NDA_BASE,
    BUYER_INITIAL_POSITION,
    SELLER_INITIAL_POSITION,
    get_disputed_clauses,
)
from src.config import OUTPUT_DIR

logger = logging.getLogger(__name__)

# Palette de couleurs redline
COLOR_AGREED  = RGBColor(0x00, 0x70, 0xC0)  # Bleu  - accord
COLOR_BUYER   = RGBColor(0xC0, 0x00, 0x00)  # Rouge - position Buyer
COLOR_SELLER  = RGBColor(0x00, 0x70, 0x00)  # Vert  - position Seller
COLOR_OPEN    = RGBColor(0xFF, 0x99, 0x00)  # Amber - clause ouverte
COLOR_BLACK   = RGBColor(0x00, 0x00, 0x00)
COLOR_GRAY    = RGBColor(0x60, 0x60, 0x60)


def _add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


class ContractWriter:
    """Genere le DOCX redline a partir du NegotiationResult."""

    def __init__(self, result):
        self.result = result
        self.doc = Document()
        self._configure_document()

    def _configure_document(self) -> None:
        section = self.doc.sections[0]
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)
        style = self.doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

    def _write_cover_page(self) -> None:
        doc = self.doc

        # Titre
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(30)
        run = title_p.add_run(NDA_BASE["title"])
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_BLACK

        # Sous-titre
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_p.add_run("NEGOTIATED DRAFT - AGL SYSTEM OUTPUT")
        sub_run.font.size = Pt(9)
        sub_run.italic = True
        sub_run.font.color.rgb = COLOR_GRAY

        doc.add_paragraph()

        # Tableau parties
        table = doc.add_table(rows=3, cols=2)
        table.style = "Table Grid"
        table.cell(0, 0).text = "BUYER"
        table.cell(0, 1).text = NDA_BASE["parties"]["buyer"]
        table.cell(1, 0).text = "SELLER"
        table.cell(1, 1).text = NDA_BASE["parties"]["seller"]
        table.cell(2, 0).text = "DATE"
        table.cell(2, 1).text = datetime.now().strftime("%B %d, %Y")

        for row in table.rows:
            for i, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
                        if i == 0:
                            run.bold = True

        doc.add_paragraph()

        # Info session
        info_p = doc.add_paragraph()
        info_run = info_p.add_run(
            "Session: " + self.result.session_id +
            "  |  Turns: " + str(self.result.total_turns) +
            "  |  Status: " + self.result.final_status +
            "  |  Convergence: " + self.result.convergence_summary.get("termination_reason", "N/A")
        )
        info_run.font.size = Pt(8)
        info_run.font.color.rgb = COLOR_GRAY
        info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

    def _write_legend(self) -> None:
        h = self.doc.add_heading("REDLINE LEGEND", level=2)
        h.paragraph_format.space_before = Pt(8)

        items = [
            (COLOR_AGREED, "Blue  - Agreed text (tentative)"),
            (COLOR_BUYER,  "Red   - Buyer position / proposed deletion"),
            (COLOR_SELLER, "Green - Seller position / proposed insertion"),
            (COLOR_OPEN,   "Amber - Open clause / still under negotiation"),
        ]
        for color, label in items:
            p = self.doc.add_paragraph(style="List Bullet")
            run = p.add_run(label)
            run.font.color.rgb = color
            run.font.size = Pt(9)

        _add_horizontal_rule(self.doc)

    def _write_clauses(self) -> None:
        self.doc.add_heading("NEGOTIATED CLAUSES", level=1)
        disputed = get_disputed_clauses()

        for clause_key, clause_data in NDA_BASE["clauses"].items():
            clause_num = clause_key.split("_")[0]

            h = self.doc.add_heading(
                "Clause " + clause_num + ": " + clause_data["title"],
                level=2
            )
            h.paragraph_format.space_before = Pt(10)

            # Texte original
            orig_p = self.doc.add_paragraph()
            orig_label = orig_p.add_run("Original: ")
            orig_label.bold = True
            orig_label.font.size = Pt(8)
            orig_label.font.color.rgb = COLOR_GRAY
            orig_text = orig_p.add_run(clause_data["text"])
            orig_text.font.size = Pt(8)
            orig_text.font.color.rgb = COLOR_GRAY
            orig_text.italic = True

            if clause_key in disputed:
                # Position Buyer
                buyer_key = "clause_" + clause_num
                buyer_pos = BUYER_INITIAL_POSITION.get(buyer_key, {})
                buyer_p = self.doc.add_paragraph()
                bl = buyer_p.add_run("BUYER: ")
                bl.bold = True
                bl.font.color.rgb = COLOR_BUYER
                bl.font.size = Pt(9)
                bt = buyer_p.add_run(buyer_pos.get("text", "N/A"))
                bt.font.color.rgb = COLOR_BUYER
                bt.font.size = Pt(9)

                # Position Seller
                seller_pos = SELLER_INITIAL_POSITION.get(buyer_key, {})
                seller_p = self.doc.add_paragraph()
                sl = seller_p.add_run("SELLER: ")
                sl.bold = True
                sl.font.color.rgb = COLOR_SELLER
                sl.font.size = Pt(9)
                st = seller_p.add_run(seller_pos.get("text", "N/A"))
                st.font.color.rgb = COLOR_SELLER
                st.font.size = Pt(9)

                # Statut
                open_p = self.doc.add_paragraph()
                open_run = open_p.add_run("STATUS: OPEN - Requires negotiation")
                open_run.font.color.rgb = COLOR_OPEN
                open_run.bold = True
                open_run.font.size = Pt(8)
            else:
                agreed_p = self.doc.add_paragraph()
                agreed_run = agreed_p.add_run("STATUS: AGREED (standard clause)")
                agreed_run.font.color.rgb = COLOR_AGREED
                agreed_run.font.size = Pt(8)

            self.doc.add_paragraph()

    def _write_transcript(self) -> None:
        self.doc.add_page_break()
        self.doc.add_heading("APPENDIX A - NEGOTIATION TRANSCRIPT", level=1)

        colors = {
            "Buyer_Counsel":  COLOR_BUYER,
            "Seller_Counsel": COLOR_SELLER,
            "Arbitrator":     COLOR_GRAY,
        }

        for entry in self.result.negotiation_log:
            speaker = entry.get("speaker", "Unknown")
            content = entry.get("content", "")
            turn    = entry.get("turn", "?")

            header_p = self.doc.add_paragraph()
            header_p.paragraph_format.space_before = Pt(6)
            header_run = header_p.add_run("[Tour " + str(turn) + "] " + speaker)
            header_run.bold = True
            header_run.font.size = Pt(9)
            header_run.font.color.rgb = colors.get(speaker, COLOR_BLACK)

            content_p = self.doc.add_paragraph()
            content_p.paragraph_format.left_indent = Inches(0.3)
            preview = content[:600] + ("..." if len(content) > 600 else "")
            content_run = content_p.add_run(preview)
            content_run.font.size = Pt(8)
            content_run.font.color.rgb = COLOR_GRAY

            _add_horizontal_rule(self.doc)

    def _write_convergence_report(self) -> None:
        self.doc.add_page_break()
        self.doc.add_heading("APPENDIX B - CONVERGENCE REPORT", level=1)

        cs = self.result.convergence_summary

        table = self.doc.add_table(rows=6, cols=2)
        table.style = "Table Grid"
        metrics = [
            ("Total Turns",          str(cs.get("total_turns", "N/A"))),
            ("Convergence Reached",  "Yes" if cs.get("convergence_reached") else "No"),
            ("Flatline Triggered",   "Yes" if cs.get("flatline_triggered") else "No"),
            ("Termination Reason",   cs.get("termination_reason", "N/A")),
            ("Average Similarity",   str(round(cs.get("avg_similarity", 0), 4))),
            ("Final Similarity",     str(round(cs.get("final_similarity", 0), 4))),
        ]
        for i, (label, value) in enumerate(metrics):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
            for run in table.cell(i, 0).paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
            for run in table.cell(i, 1).paragraphs[0].runs:
                run.font.size = Pt(9)

    def generate(self, output_path: Optional[str] = None) -> str:
        """Genere le DOCX complet et retourne le chemin du fichier."""
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        if output_path is None:
            output_path = os.path.join(
                OUTPUT_DIR,
                "negotiated_nda_" + self.result.session_id + ".docx"
            )

        self._write_cover_page()
        self._write_legend()
        self._write_clauses()
        self._write_transcript()
        self._write_convergence_report()

        self.doc.save(output_path)
        logger.info("DOCX genere : " + output_path)
        return output_path

